from docx import Document
from docx.shared import Pt

def replace_placeholders_in_text(text, context):
    for key, value in context.items():
        text = text.replace(f"{{{{{key}}}}}", str(value))
    return text

def process_paragraphs(paragraphs, context):
    for paragraph in paragraphs:
        original_text = ''.join(run.text for run in paragraph.runs)
        replaced_text = replace_placeholders_in_text(original_text, context)

        paragraph.clear()
        run = paragraph.add_run(replaced_text)

        # If it's a heading, apply bold + underline + size 14
        if "SALE AGREEMENT FOR MOTOR VEHICLE" in replaced_text or "CHASSIS NO." in replaced_text:
            run.bold = True
            run.underline = True
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(11)

def process_tables(tables, context):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, context)

def fill_template(template_path, output_path, context):
    doc = Document(template_path)

    # Apply processing
    process_paragraphs(doc.paragraphs, context)
    process_tables(doc.tables, context)

    doc.save(output_path)
    print(f"✅ Document saved to: {output_path}")

# --- Your dynamic context ---
context = {
    "client_name": "James Gacheru",
    "vehicle_reg": "KDA 123X",
    "price": "1,500,000",
    "sale_date": "14th July 2025",
    "amount_words": "One Million Five Hundred Thousand",
    "amount_no": "1,500,000",
    "yom": "2016",
    "vehicle_make": "Toyota Hilux",
    "vehicle_color": "Black",
    "chasis_no": "AIEBFIWFP",
    "engine_no": "XYZ987654321",
    "client_id": "35867077",
    "email_id": "gacherujay@gmail.com",
    "contact_no": "0714404065",
    "postal_address": "797-00100 Nairobi ",
}

template_path = "templates/sales_agreement.docx"
output_path = "generated_docs/sales_agreement_update.docx"

fill_template(template_path, output_path, context)
