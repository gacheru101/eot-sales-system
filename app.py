# === IMPORTS ===

from docx import Document
from docx.shared import Pt
from num2words import num2words
import sqlite3
import pandas as pd
import os
from database import init_db
from flask import Flask, render_template, request, send_file, redirect, url_for, flash
from flask import session
from werkzeug.security import check_password_hash
from werkzeug.security import generate_password_hash

from functools import wraps
from flask import session, redirect, url_for, flash

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'username' not in session:
            flash("❌ Please log in to continue.")
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if session.get("role") != "admin":
            flash("🚫 Admin access required.")
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return decorated_function



# === INIT APP AND DB ===
app = Flask(__name__)
app.secret_key = "super_secret_key_2025"
init_db()

print("🔥 Flask is running from:", os.path.abspath(__file__))

# === LOAD VEHICLE DATA ===
stock_df = pd.read_excel("vehicles.xlsx", sheet_name="Stock List")
sold_df = pd.read_excel("vehicles.xlsx", sheet_name="Sold")

stock_df["Reg No"] = stock_df["Reg No"].astype(str).str.upper()
sold_df["Reg No"] = sold_df["Reg No"].astype(str).str.upper()

vehicle_lookup = {
    row["Reg No"]: {
        "chassis": row["Chassis/Frame Number"],
        "engine": row["Engine Number"]
    }
    for _, row in stock_df.iterrows()
}

# === DOCUMENT FUNCTIONS ===
def replace_placeholders_in_text(text, context):
    for key, value in context.items():
        text = text.replace(f"{{{{{key}}}}}", str(value))
    return text

def process_paragraphs(paragraphs, context):
    for paragraph in paragraphs:
        original_text = ''.join(run.text for run in paragraph.runs)
        replaced_text = replace_placeholders_in_text(original_text, context)

        paragraph.clear()
        run = paragraph.add_run(replaced_text)

        if "SALE AGREEMENT FOR MOTOR VEHICLE" in replaced_text or "CHASSIS NO." in replaced_text:
            run.bold = True
            run.underline = True
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(11)

def process_tables(tables, context):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, context)

def fill_template(template_path, output_path, context):
    doc = Document(template_path)
    process_paragraphs(doc.paragraphs, context)
    process_tables(doc.tables, context)
    doc.save(output_path)
def fill_template_bold(template_path, output_path, context):
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        original = ''.join(run.text for run in paragraph.runs)
        replaced = replace_placeholders_in_text(original, context)
        paragraph.clear()
        run = paragraph.add_run(replaced)
        run.font.size = Pt(11)

        # Bold top-left reference details
        if any(keyword in replaced.lower() for keyword in ["ref", "to:", "date:", "attention", "our ref"]):
            run.bold = True

        # Bold signature or sign-off section
        if any(keyword in replaced.lower() for keyword in ["yours faithfully", "alexander", "bosire", "regards"]):
            run.bold = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    original = ''.join(run.text for run in p.runs)
                    replaced = replace_placeholders_in_text(original, context)
                    p.clear()
                    run = p.add_run(replaced)
                    run.font.size = Pt(11)

    doc.save(output_path)


# === ROUTES ===
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        conn = sqlite3.connect("booking.db")
        c = conn.cursor()
        c.execute("SELECT id, username, password, role FROM users WHERE username = ?", (username,))
        user = c.fetchone()
        conn.close()

        if user and check_password_hash(user[2], password):
            session["user_id"] = user[0]
            session["username"] = user[1]
            session["role"] = user[3]
            flash(f"✅ Welcome {user[1]}!", "success")
            return redirect(url_for("index"))
        else:
            flash("❌ Invalid username or password.", "danger")
            return redirect(url_for("login"))

    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    flash("🔒 You have been logged out.")
    return redirect(url_for("login"))



@app.route("/upload", methods=["GET", "POST"])
@login_required
def upload_excel():
    if request.method == "POST":
        file = request.files["excel_file"]
        if file.filename.endswith(".xlsx"):
            df_new = pd.read_excel(file, sheet_name="Stock List")
            df_new["Reg No"] = df_new["Reg No"].astype(str).str.upper()

            # Load current stock
            global stock_df
            stock_df = pd.concat([stock_df, df_new]).drop_duplicates(subset="Reg No", keep="last").reset_index(drop=True)

            # Update vehicle_lookup
            global vehicle_lookup
            vehicle_lookup = {
                row["Reg No"]: {
                    "chassis": row["Chassis/Frame Number"],
                    "engine": row["Engine Number"]
                }
                for _, row in stock_df.iterrows()
            }

            flash("✅ New vehicles uploaded and merged successfully.")
            return redirect(url_for("index"))
        else:
            flash("❌ Please upload a valid .xlsx file.")
            return redirect(url_for("upload_excel"))
    
    return render_template("upload.html")

@app.route("/")
def index():
    conn = sqlite3.connect("booking.db")
    c = conn.cursor()
    c.execute("""
        SELECT 
            b.vehicle_reg,
            b.price,
            IFNULL(SUM(p.amount), 0) as total_paid,
            b.price - IFNULL(SUM(p.amount), 0) as balance
        FROM bookings b
        LEFT JOIN payments p ON b.vehicle_reg = p.vehicle_reg
        GROUP BY b.vehicle_reg
        ORDER BY b.vehicle_reg
    """)
    summary_rows = c.fetchall()
    conn.close()
    return render_template("form.html", vehicle_lookup=vehicle_lookup, summary_rows=summary_rows)


@app.route("/record-payment", methods=["GET"])
@login_required
def payment_form():
    conn = sqlite3.connect("booking.db")
    c = conn.cursor()
    c.execute("SELECT vehicle_reg FROM bookings WHERE status = 'booked'")
    booked = [row[0] for row in c.fetchall()]
    conn.close()

    return render_template("payment_form.html", booked_vehicles=booked)

@app.route("/bookings")
@login_required
def view_bookings():
    conn = sqlite3.connect("booking.db")
    c = conn.cursor()
    c.execute("SELECT vehicle_reg, client_name, client_id, sale_date, price, status FROM bookings ORDER BY vehicle_reg")
    bookings = c.fetchall()
    conn.close()
    return render_template("bookings.html", bookings=bookings)



@app.route("/record-payment", methods=["POST"])
@login_required
def record_payment():
    vehicle_reg = request.form["vehicle_reg"].upper()
    amount = int(request.form["amount"])
    reference = request.form["reference"]
    date = request.form["date"]

    conn = sqlite3.connect("booking.db")
    c = conn.cursor()

    # ✅ Check if price exists
    c.execute("SELECT price FROM bookings WHERE UPPER(vehicle_reg) = ?", (vehicle_reg,))
    result = c.fetchone()

    if not result or result[0] is None:
        conn.close()
        flash(f"❌ Vehicle {vehicle_reg} does not have a valid price recorded in bookings. Please rebook or update the price in DB.")
        return redirect(url_for("payment_form"))

    expected_price = result[0]

    # ✅ Record payment
    c.execute("""
        INSERT INTO payments (vehicle_reg, amount, reference, date)
        VALUES (?, ?, ?, ?)
    """, (vehicle_reg, amount, reference, date))

    # ✅ Get total paid
    c.execute("SELECT SUM(amount) FROM payments WHERE vehicle_reg = ?", (vehicle_reg,))
    total_paid = c.fetchone()[0] or 0

    # ✅ Update status to sold if fully paid
    if total_paid >= expected_price:
        c.execute("UPDATE bookings SET status = 'sold' WHERE vehicle_reg = ?", (vehicle_reg,))

    conn.commit()
    conn.close()

    # ✅ Message
    overpayment = total_paid - expected_price if total_paid > expected_price else 0
    message = f"✅ Payment of KES {amount:,} recorded for {vehicle_reg}. Total Paid: KES {total_paid:,}."
    if overpayment > 0:
        message += f" (Overpayment: KES {overpayment:,})"

    flash(message)
    return redirect(url_for("payment_form"))




@app.route("/payments")
@login_required
def view_payments():
    conn = sqlite3.connect("booking.db")
    c = conn.cursor()

    c.execute("""
        SELECT vehicle_reg, date, reference, amount
        FROM payments
        ORDER BY vehicle_reg, date
    """)
    rows = c.fetchall()
    conn.close()

    return render_template("payments.html", payments=rows)



import zipfile

@app.route("/generate", methods=["POST"])
@login_required
def generate():
    data = request.form.to_dict()
    vehicle_reg = data.get("vehicle_reg").upper()
    price_str = data.get("price").replace(",", "")
    price_int = int(price_str)

    vehicle_info = vehicle_lookup.get(vehicle_reg)
    if not vehicle_info:
        return f"<h3 style='color:red;'>❌ Could not find vehicle info for {vehicle_reg}.</h3>"

    chasis_no = vehicle_info["chassis"]
    engine_no = vehicle_info["engine"]
    client_name = data.get("client_name")
    client_id = data.get("client_id")
    sale_date = data.get("sale_date")

    # === Check if already booked ===
    conn = sqlite3.connect("booking.db")
    c = conn.cursor()
    c.execute("SELECT * FROM bookings WHERE vehicle_reg = ?", (vehicle_reg,))
    existing = c.fetchone()
    if existing:
        conn.close()
        return f"<h3 style='color:red;'>❌ Vehicle {vehicle_reg} is already BOOKED.</h3>"

    # === Save new booking ===
    c.execute("""
        INSERT INTO bookings (vehicle_reg, client_name, client_id, sale_date, price)
        VALUES (?, ?, ?, ?, ?)
    """, (vehicle_reg, client_name, client_id, sale_date, price_int))
    conn.commit()
    conn.close()

    amount_words = num2words(price_int, to='cardinal', lang='en').title()
    date_str = data.get("sale_date")
    make_model = f"{data.get('vehicle_make')} {data.get('vehicle_color')}"

    # === Shared context ===
    context = {
        "client_name": client_name,
        "vehicle_reg": vehicle_reg,
        "price": data.get("price"),
        "sale_date": sale_date,
        "amount_words": amount_words,
        "amount_no": f"{price_int:,}",
        "yom": data.get("yom"),
        "vehicle_make": data.get("vehicle_make"),
        "vehicle_color": data.get("vehicle_color"),
        "chasis_no": chasis_no,
        "engine_no": engine_no,
        "client_id": client_id,
        "email_id": data.get("email_id"),
        "contact_no": data.get("contact_no"),
        "postal_address": data.get("postal_address"),
        "make_model": make_model,
        "date": date_str,
        "reg": vehicle_reg,
        "director1": "Alexander Wambua",
        "director2": "Bosire Bogonko",
    }

    # === Generate all 4 docs ===
    sales_doc = f"generated_docs/{vehicle_reg}_agreement.docx"
    ntsa_doc = f"generated_docs/{vehicle_reg}_ntsa_letter.docx"
    undertaking_doc = f"generated_docs/{vehicle_reg}_undertaking.docx"
    release_doc = f"generated_docs/{vehicle_reg}_release_letter.docx"

    fill_template("templates/sales_agreement.docx", sales_doc, context)  # leave this as is
    fill_template_bold("templates/ntsa_letter.docx", ntsa_doc, context)  # use bold function
    fill_template_bold("templates/undertaking_letter.docx", undertaking_doc, context)
    fill_template_bold("templates/release_letter.docx", release_doc, context)

    # === Zip them ===
    zip_path = f"generated_docs/{vehicle_reg}_documents.zip"
    with zipfile.ZipFile(zip_path, "w") as zipf:
        zipf.write(sales_doc, os.path.basename(sales_doc))
        zipf.write(ntsa_doc, os.path.basename(ntsa_doc))
        zipf.write(undertaking_doc, os.path.basename(undertaking_doc))
        zipf.write(release_doc, os.path.basename(release_doc))

    return send_file(zip_path, as_attachment=True)




@app.route("/generate-release", methods=["POST"])
@login_required
def generate_release():
    data = request.form.to_dict()
    vehicle_reg  = data["vehicle_reg"].upper()
    client_name  = data["client_name"]
    client_id    = data["client_id"]
    contact_no   = data["contact_no"]
    date         = data["date"]
    chasis_no    = vehicle_lookup[vehicle_reg]["chassis"]
    engine_no    = vehicle_lookup[vehicle_reg]["engine"]
    make_model   = data["make_model"]
    # Directors' names could also come from config or from form:
    director1    = "Alexander Wambua"
    director2    = "Bosire Bogonko"

    context = {
      "reg": vehicle_reg,
      "client_name": client_name,
      "client_id": client_id,
      "contact_no": contact_no,
      "date": date,
      "chasis_no": chasis_no,
      "engine_no": engine_no,
      "make_model": make_model,
      "director1": director1,
      "director2": director2,
    }

    template_path = "templates/release_letter.docx"
    output_path   = f"generated_docs/{vehicle_reg}_release_letter.docx"
    fill_template_bold(template_path, output_path, context)

    return send_file(output_path, as_attachment=True)

@app.route("/payments-summary")
@login_required
def payments_summary():
    conn = sqlite3.connect("booking.db")
    c = conn.cursor()

    c.execute("""
        SELECT 
            b.vehicle_reg,
            b.price,
            IFNULL(SUM(p.amount), 0) as total_paid,
            b.price - IFNULL(SUM(p.amount), 0) as balance
        FROM bookings b
        LEFT JOIN payments p ON b.vehicle_reg = p.vehicle_reg
        GROUP BY b.vehicle_reg
        ORDER BY b.vehicle_reg
    """)
    rows = c.fetchall()
    conn.close()

    return render_template("payment_summary.html", rows=rows)
from flask import send_file
from datetime import datetime

# ---- NTSA Letter ----
@app.route("/generate-ntsa", methods=["POST"])
@login_required
def generate_ntsa():
    data = request.form.to_dict()
    reg         = data["vehicle_reg"].upper()
    client_name = data["client_name"]
    client_id   = data["client_id"]
    contact_no  = data["contact_no"]
    postal_addr = data["postal_address"]
    make_model  = data["make_model"]
    yom         = data["yom"]
    price       = data["price"]
    date_str    = data["date"] or datetime.today().strftime("%d %B %Y")

    ch_info = vehicle_lookup.get(reg, {})
    chasis_no = ch_info.get("chassis", "")
    engine_no = ch_info.get("engine", "")

    context = {
      "date":           date_str,
      "reg":            reg,
      "client_name":    client_name,
      "client_id":      client_id,
      "contact_no":     contact_no,
      "postal_address": postal_addr,
      "make_model":     make_model,
      "yom":            yom,
      "chasis_no":      chasis_no,
      "engine_no":      engine_no,
    }

    template_path = "templates/ntsa_letter.docx"
    output_path   = f"generated_docs/{reg}_ntsa_letter.docx"
    fill_template_bold(template_path, output_path, context)
    return send_file(output_path, as_attachment=True)


# ---- Letter of Undertaking ----
@app.route("/generate-undertaking", methods=["POST"])
@login_required
def generate_undertaking():
    data = request.form.to_dict()
    reg         = data["vehicle_reg"].upper()
    client_name = data["client_name"]
    client_id   = data["client_id"]
    contact_no  = data["contact_no"]
    make_model  = data["make_model"]
    yom         = data["yom"]
    price       = data["price"]
    date_str    = data["date"] or datetime.today().strftime("%d %B %Y")

    ch_info = vehicle_lookup.get(reg, {})
    chasis_no = ch_info.get("chassis", "")
    engine_no = ch_info.get("engine", "")

    context = {
      "date":        date_str,
      "reg":         reg,
      "client_name": client_name,
      "client_id":   client_id,
      "contact_no":  contact_no,
      "make_model":  make_model,
      "yom":         yom,
      "price":       price,
      "chasis_no":   chasis_no,
      "engine_no":   engine_no,
    }

    template_path = "templates/undertaking_letter.docx"
    output_path   = f"generated_docs/{reg}_undertaking.docx"
    fill_template_bold(template_path, output_path, context)
    return send_file(output_path, as_attachment=True)






@app.route("/download-report")
@login_required
def download_report():
    # Load stock list from Excel
    stock_df = pd.read_excel("vehicles.xlsx", sheet_name="Stock List")

    # Connect to DB and load booking and payment data
    conn = sqlite3.connect("booking.db")

    # Booked = all booked vehicles
    booked_df = pd.read_sql_query("SELECT * FROM bookings WHERE status = 'booked'", conn)

    # Sold = all paid-off vehicles
    sold_df = pd.read_sql_query("SELECT * FROM bookings WHERE status = 'sold'", conn)

    conn.close()

    # Save to Excel with 3 sheets
    report_path = "Vehicle_Report.xlsx"
    with pd.ExcelWriter(report_path, engine="openpyxl") as writer:
        stock_df.to_excel(writer, sheet_name="Stock List", index=False)
        booked_df.to_excel(writer, sheet_name="Booked Units", index=False)
        sold_df.to_excel(writer, sheet_name="Sold Units", index=False)

    return send_file(report_path, as_attachment=True)
@app.route("/missing-prices", methods=["GET", "POST"])
@login_required
def fix_missing_prices():
    conn = sqlite3.connect("booking.db")
    c = conn.cursor()

    if request.method == "POST":
        vehicle_reg = request.form["vehicle_reg"].upper()
        price = int(request.form["price"])

        # Update price in bookings
        c.execute("UPDATE bookings SET price = ? WHERE vehicle_reg = ?", (price, vehicle_reg))
        conn.commit()

    # Show all bookings where price is NULL
    c.execute("SELECT vehicle_reg, client_name, sale_date FROM bookings WHERE price IS NULL OR price = ''")
    rows = c.fetchall()
    conn.close()

    return render_template("missing_prices.html", bookings=rows)
@app.route("/admin/bookings")
@admin_required
def view_all_bookings():
    conn = sqlite3.connect("booking.db")
    c = conn.cursor()
    c.execute("SELECT vehicle_reg, client_name, sale_date, price FROM bookings")
    bookings = c.fetchall()
    conn.close()
    return render_template("admin_bookings.html", bookings=bookings)

@app.route("/admin/bookings/delete/<vehicle_reg>")
@admin_required
def delete_booking(vehicle_reg):
    conn = sqlite3.connect("booking.db")
    c = conn.cursor()
    c.execute("DELETE FROM bookings WHERE vehicle_reg = ?", (vehicle_reg,))
    conn.commit()
    conn.close()
    return f"<h3 style='color:red;'>❌ Booking for {vehicle_reg} deleted.</h3><a href='/admin/bookings'>🔙 Back to bookings</a>"
@app.route("/admin/payments")
@admin_required
def view_all_payments():
    conn = sqlite3.connect("booking.db")
    c = conn.cursor()
    c.execute("SELECT id, vehicle_reg, amount, reference, date FROM payments ORDER BY date DESC")
    payments = c.fetchall()
    conn.close()
    return render_template("admin_payments.html", payments=payments)

@app.route("/admin/payments/delete/<int:payment_id>")
@admin_required
def delete_payment(payment_id):
    conn = sqlite3.connect("booking.db")
    c = conn.cursor()
    c.execute("DELETE FROM payments WHERE id = ?", (payment_id,))
    conn.commit()
    conn.close()
    return f"<h3 style='color:red;'>❌ Payment deleted.</h3><a href='/admin/payments'>🔙 Back to payments</a>"



@app.route("/register", methods=["GET", "POST"])
@admin_required
def register():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        role = request.form["role"]

        hashed_pw = generate_password_hash(password)

        conn = sqlite3.connect("booking.db")
        c = conn.cursor()
        try:
            c.execute("INSERT INTO users (username, password, role) VALUES (?, ?, ?)", 
                      (username, hashed_pw, role))
            conn.commit()
            flash(f"✅ User '{username}' created successfully.", "success")
        except sqlite3.IntegrityError:
            flash(f"❌ Username '{username}' already exists.", "danger")
        finally:
            conn.close()
        return redirect(url_for("register"))

    return render_template("register.html")




# === RUN FLASK ===
if __name__ == "__main__":
    app.run(debug=True)

